from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docxtpl import DocxTemplate
from aspose.words import Document as AsposeDocument
from aspose.words import License
from sys import platform
from docx2pdf import convert


from .settings import TOC_FONT, TITLE_FONT, CONTENT_FONT


class DocumentGenerator:
    def __init__(self, book) -> None:
        self.book = book

        try:
            # Set the Aspose license for the document generator
            self.license = License()
            self.license.set_license("aspose.lic")
        except Exception as e:
            print(f"Error setting Aspose license: {e}")

        self.template = None
        self.document = None

    def generate_cover_page(self) -> None:
        self.template = DocxTemplate("templates/literature.docx")

        # Replace the placeholders in the template
        context = {
            "AUTHOR": self.book["author"].upper(),
            "TITLE": self.book["title"],
        }

        # Set cover picture
        self.template.replace_pic("cover.png", self.book["cover"])

        # Render the template and save it
        self.template.render(context)
        self.template.save(f"media/docs/{self.book['id']}.docx")

    def generate_pdf(self) -> None:
        try:
            # Convert the document to PDF using Aspose (cross-platform)
            doc = AsposeDocument(f"media/docs/{self.book['id']}.docx")
            doc.save(f"media/pdfs/{self.book['id']}.pdf")

        # If Aspose fails, try using more platform-specific libraries
        except Exception as e:
            print(f"Error converting to PDF using Aspose: {e}")

            if platform.startswith("linux"):
                try:
                    from subprocess import call

                    # Convert the document to PDF using LibreOffice (Linux only)
                    call(
                        [
                            "libreoffice",
                            "--headless",
                            "--convert-to",
                            "pdf",
                            "--outdir",
                            "media/pdfs",
                            f"media/docs/{self.book['id']}.docx",
                        ]
                    )
                except Exception as e:
                    print(f"Error converting to PDF using LibreOffice: {e}")

            else:
                try:
                    # Convert the document to PDF using docx2pdf (Windows/MacOS only)
                    convert(
                        f"media/docs/{self.book['id']}.docx",
                        f"media/pdfs/{self.book['id']}.pdf",
                    )
                except Exception as e:
                    print(f"Error converting to PDF using docx2pdf: {e}")

    def generate_document(self) -> None:
        self.document = Document(f"media/docs/{self.book['id']}.docx")

        # Add table of contents
        self.add_table_of_contents()

        # Add content
        self.add_content()

        # Save the document
        self.document.save(f"media/docs/{self.book['id']}.docx")

    def add_table_of_contents(self) -> None:
        self.document.add_page_break()

        h = self.document.add_heading("Table of Contents", 0)
        h.runs[0].font.name = TOC_FONT

        self.document.add_paragraph()

        for chapter in self.book.get("table_of_contents"):

            h = self.document.add_heading(chapter.get("chapter"), 1)
            h.runs[0].font.name = TOC_FONT
            h.runs[0].font.size = Pt(20)

            for subsection in chapter.get("subsections"):

                section_number, section_title = subsection.split(" ", 1)

                h = self.document.add_heading(level=2)

                r = h.add_run(section_number + " ")
                r.font.name = TITLE_FONT
                r.font.size = Pt(18)
                r.bold = True

                r = h.add_run(section_title)
                r.font.name = TITLE_FONT
                r.font.size = Pt(18)

            self.document.add_paragraph()

    def add_content(self) -> None:
        for chapter in self.book["content"]:
            self.document.add_page_break()

            # Add chapter title to the document in the center of a blank page
            for _ in range(13):
                self.document.add_paragraph()

            chapter_title = self.document.add_paragraph()

            ch = chapter.get("chapter")
            if chapter.get("chapter").lower().startswith(
                "chapter"
            ) and ":" in chapter.get("chapter"):
                ch = chapter.get("chapter").split(":", 1)[1].strip()

            r = chapter_title.add_run(ch)
            r.bold = True
            r.font.size = Pt(30)
            r.font.name = TITLE_FONT

            chapter_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            for subsection in chapter["subsections"]:

                self.document.add_page_break()

                h = self.document.add_heading(subsection["subsection"], 1)
                h.runs[0].font.name = TITLE_FONT
                h.runs[0].font.size = Pt(24)

                first = True
                for paragraph in subsection["paragraphs"]:

                    if (
                        paragraph.lower() == subsection["subsection"].lower()
                        or paragraph.lower()
                        == subsection["subsection"].split(" ", 1)[1].strip().lower()
                    ):
                        continue

                    p = self.document.add_paragraph(paragraph)

                    if first:
                        first = False
                        p.paragraph_format.space_before = Pt(12)

                    p.paragraph_format.space_after = Pt(12)
                    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    p.runs[0].font.name = CONTENT_FONT
                    p.runs[0].font.size = Pt(16)
